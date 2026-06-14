"""Fill the CV_Project_Report_Template.docx with real project content + figures.

Run AFTER train.py and make_figures.py:  python make_report.py
Output: CV_Project_Report_FILLED.docx   (template itself is left untouched)

Identifying details (name, course, etc.) are read from report_meta.json if it
exists; otherwise clear [PLACEHOLDERS] are inserted for you to fill.
"""
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

TEMPLATE = config.ROOT / "CV_Project_Report_Template.docx"
OUTPUT = config.ROOT / "CV_Project_Report_FILLED.docx"
RESULTS = config.ROOT / "results"
labels = config.DISEASE_LABELS

# ---------------------------------------------------------------- load real data
metrics = json.loads((config.ARTIFACTS_DIR / "metrics.json").read_text())
history = json.loads((config.ARTIFACTS_DIR / "history.json").read_text())
df = pd.read_csv(config.SPLITS_CSV)

counts = {l: int(df[l].sum()) for l in labels}
counts_sorted = sorted(counts.items(), key=lambda kv: -kv[1])
n_total = len(df)
n_train = int((df["split"] == "train").sum())
n_val = int((df["split"] == "val").sum())
n_test = int((df["split"] == "test").sum())
best_ep = max(history, key=lambda h: h["val_macro_auc"])["epoch"]

auc = metrics["test_macro_auc"]
best_val_auc = metrics["best_val_macro_auc"]
macro_f1 = metrics["test_macro_f1"]
micro_f1 = metrics["test_micro_f1"]
macro_p = metrics["test_macro_precision"]
macro_r = metrics["test_macro_recall"]
per_label_acc = metrics["test_per_label_accuracy"]
subset_acc = metrics["test_subset_accuracy"]
pc = metrics["per_class"]
sup_sum = sum(pc[l]["support"] for l in labels)
baseline_acc = 1 - sup_sum / (config.NUM_CLASSES * n_test)   # all-negative predictor
top_classes = sorted(pc, key=lambda n: -pc[n]["auc"])

# identifying info
meta_path = config.ROOT / "report_meta.json"
meta = json.loads(meta_path.read_text()) if meta_path.exists() else {}
AUTHOR = meta.get("author", "[Your Name]")
SID = meta.get("student_id", "[Student ID]")
COURSE = meta.get("course", "[Course Title & Code]")
INSTRUCTOR = meta.get("instructor", "[Instructor Name]")
UNIV = meta.get("university", "[University / Department]")
DATE = meta.get("date", "14 June 2026")

doc = Document(TEMPLATE)

# ================================================================ helpers
def key_of(text):
    """Leading section number e.g. '4.2.1' -> '4.2.1'; else None."""
    tok = text.strip().split(" ")[0] if text.strip() else ""
    return tok if tok and tok[0].isdigit() else None


def add_after(anchor_el, items):
    """Insert a list of content items right after a body element (in order)."""
    cur = anchor_el
    for it in items:
        t = it[0]
        if t == "p":
            p = doc.add_paragraph(it[1])
        elif t == "b":
            p = doc.add_paragraph(it[1], style="List Paragraph")
        elif t == "cap":
            p = doc.add_paragraph()
            r = p.add_run(it[1]); r.italic = True; r.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "img":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(RESULTS / it[1]), width=Inches(it[2]))
        cur.addnext(p._p)
        cur = p._p
    return cur


def fill_table(tbl, values):
    """values: list of cell texts for the LAST column, one per row (skip header rows handled by caller)."""
    for row, val in values:
        tbl.rows[row].cells[-1].text = val


# ================================================================ content
C = {}

C["Abstract"] = [("p",
    f"Chest radiography is the most common medical imaging examination worldwide, "
    f"but expert interpretation is time-consuming and unevenly available. This project "
    f"develops an automated system that performs multi-label classification of chest "
    f"X-rays, simultaneously predicting the presence of 14 thoracic conditions. We use "
    f"the publicly available NIH ChestX-ray14 sample dataset ({n_total:,} frontal X-ray "
    f"images with weakly-supervised disease labels mined from radiology reports). Because "
    f"the work targets a low-resource, CPU-only environment, we adopt a transfer-learning "
    f"strategy: a MobileNetV2 network pretrained on ImageNet is used as a frozen feature "
    f"extractor, its 1280-dimensional features are cached, and a compact fully-connected "
    f"classifier head is trained on top using a sigmoid output and class-weighted binary "
    f"cross-entropy to address severe class imbalance. The data are split by patient "
    f"({n_train}/{n_val}/{n_test} train/validation/test) to avoid leakage. On the held-out "
    f"test set the model attains a macro-averaged ROC-AUC of {auc:.2f}, with the strongest "
    f"performance on {top_classes[0]} (AUC {pc[top_classes[0]]['auc']:.2f}), "
    f"{top_classes[1]} ({pc[top_classes[1]]['auc']:.2f}) and "
    f"{top_classes[2]} ({pc[top_classes[2]]['auc']:.2f}). Training curves reveal early "
    f"overfitting, with the best validation performance reached at epoch {best_ep}. We "
    f"conclude that frozen-feature transfer learning is a viable, reproducible baseline "
    f"for multi-label chest X-ray classification without specialised hardware, while "
    f"class imbalance and limited data remain the principal obstacles to higher accuracy.")]

C["1.1"] = [("p",
    "Chest X-ray (CXR) imaging is a cornerstone of modern diagnostic medicine because it "
    "is fast, inexpensive and low-dose. However, the global volume of radiographs far "
    "exceeds the capacity of trained radiologists, and interpretation is subjective and "
    "error-prone, particularly in under-resourced settings. Computer-aided diagnosis "
    "powered by deep learning has therefore become an active research area [1], [2]. The "
    "release of large, labelled datasets such as ChestX-ray14 [1] and CheXpert [3] has "
    "enabled convolutional neural networks (CNNs) to reach, and in some narrow tasks "
    "exceed, radiologist-level performance [2]."),
    ("p",
    "A key characteristic of real chest radiographs is that a single image frequently "
    "exhibits several abnormalities at once (for example, effusion together with "
    "cardiomegaly). The task is thus inherently multi-label rather than single-class, "
    "which motivates models that output an independent probability for each condition.")]

C["1.2"] = [("p",
    "This project builds a system that takes a single frontal chest X-ray as input and "
    "outputs the probability of each of 14 thoracic findings, allowing multiple findings "
    "to be reported simultaneously. Existing high-performing solutions typically rely on "
    "fine-tuning very deep networks end-to-end using GPU hardware and the full "
    "112,000-image dataset [2], [3], which is impractical for students and practitioners "
    "without such resources. The problem addressed here is to obtain a credible "
    "multi-label CXR classifier under a strict CPU-only, limited-data constraint.")]

C["1.3"] = [
    ("b", "To develop a multi-label classifier that predicts 14 thoracic conditions from a single chest X-ray."),
    ("b", "To apply transfer learning with a frozen pretrained CNN backbone so that the system trains on a CPU."),
    ("b", "To address class imbalance using class-weighted loss and threshold-independent evaluation."),
    ("b", f"To evaluate the model with ROC-AUC and F1-score, targeting a macro ROC-AUC of at least 0.65 on the ChestX-ray14 sample test set."),
    ("b", "To deliver a reproducible pipeline and an interactive demonstration application."),
]

C["1.4"] = [("p",
    "The system is intended as an educational decision-support prototype, not a certified "
    "diagnostic device. It is trained only on the 5,606-image ChestX-ray14 sample, so its "
    "performance does not represent what is achievable on the full dataset. Labels are "
    "weakly supervised (text-mined from reports) and therefore noisy. The model uses only "
    "frontal images, performs classification but not localisation (no bounding boxes), "
    "and the backbone is frozen rather than fine-tuned. Clinical deployment would require "
    "far larger, curated data, prospective validation and regulatory approval.")]

# Literature review themes (rename headings + content)
THEME_TITLES = {
    "2.1": "2.1 CNN-Based Chest X-ray Classification",
    "2.2": "2.2 Transfer Learning and Pretrained Backbones",
    "2.3": "2.3 Datasets and Benchmarks",
}
C["2.1"] = [("p",
    "Wang et al. [1] introduced ChestX-ray8/14 and established CNN baselines for "
    "weakly-supervised classification of common thorax diseases, framing the task as "
    "multi-label. Rajpurkar et al. [2] proposed CheXNet, a 121-layer DenseNet trained on "
    "the full dataset, reporting radiologist-level pneumonia detection and high AUCs "
    "across all 14 findings. Irvin et al. [3] extended this line with CheXpert, a larger "
    "dataset with explicit uncertainty labels, and benchmarked several CNN architectures. "
    "These works consistently report ROC-AUC as the primary metric because of severe class "
    "imbalance, but they depend on end-to-end GPU training over very large datasets.")]
C["2.2"] = [("p",
    "Transfer learning reuses representations learned on large source datasets such as "
    "ImageNet [7], [12] for data-scarce target tasks [8], [9]. Architectures including "
    "ResNet [6] and DenseNet [4] are popular backbones, while MobileNetV2 [5] offers a "
    "lightweight inverted-residual design well suited to constrained hardware. A common "
    "and efficient variant freezes the convolutional base and trains only a new classifier "
    "head, which dramatically reduces computation — the approach adopted in this project.")]
C["2.3"] = [("p",
    "Public CXR benchmarks include ChestX-ray14 [1] and CheXpert [3]; the survey by Çallı "
    "et al. [10] reviews datasets, label-noise issues and methodological pitfalls in deep "
    "learning for chest radiography. A recurring theme is that headline accuracy is "
    "misleading on imbalanced data, and that patient-level data splitting is essential to "
    "avoid over-optimistic results. Implementations in this area are dominated by deep "
    "learning frameworks such as PyTorch [11].")]
C["2.4"] = [("p",
    "Most prior work assumes abundant GPU compute, end-to-end fine-tuning and the complete "
    "112,000-image dataset. Comparatively little attention is paid to whether a useful "
    "multi-label classifier can be obtained under strict low-resource constraints. This "
    "project addresses that gap by quantifying the performance of a frozen-feature "
    "transfer-learning pipeline trained entirely on a CPU using only the 5,606-image "
    "sample, and by documenting a fully reproducible workflow.")]

C["3.1"] = [("p",
    "The dataset is the Random Sample of the NIH Chest X-ray dataset, made available on "
    "Kaggle (https://www.kaggle.com/datasets/nih-chest-xrays/sample) and derived from the "
    "ChestX-ray14 database released by the U.S. National Institutes of Health [1]. It "
    "contains 5,606 frontal-view radiographs with multi-label disease annotations that "
    "were text-mined from the associated radiology reports.")]

C["3.3"] = [("p",
    "Figure 1 shows representative chest X-rays from the dataset, each annotated with its "
    "ground-truth findings. The examples illustrate the grayscale nature of the images and "
    "the fact that many radiographs carry several labels simultaneously."),
    ("img", "fig1_sample_grid.png", 6.3)]

C["3.4"] = [("p",
    f"The dataset is severely class-imbalanced. Across the sample, Infiltration is the "
    f"most frequent finding ({counts['Infiltration']} positives) while Hernia is extremely "
    f"rare ({counts['Hernia']} positives); the majority of label slots are negative. Labels "
    f"are also noisy because they are automatically extracted from free text rather than "
    f"verified by radiologists, and the images are single-channel grayscale at varying "
    f"effective resolution. These characteristics directly motivated three design choices: "
    f"(i) resizing and 3-channel replication so images match the pretrained backbone; "
    f"(ii) class-weighted loss so rare findings are not ignored; and (iii) evaluation with "
    f"ROC-AUC and F1 rather than plain accuracy.")]

C["4.1"] = [("p",
    "The system follows a five-stage pipeline (Figure 2): (1) an input chest X-ray is "
    "(2) preprocessed, (3) passed through a frozen MobileNetV2 backbone to produce a "
    "1280-dimensional feature vector that is cached to disk, (4) classified by a small "
    "trainable head, and (5) converted to 14 independent probabilities via a sigmoid "
    "activation. Decoupling feature extraction from training is what makes the workflow "
    "feasible on a CPU: the expensive convolutional pass is computed once, after which the "
    "head trains in seconds."),
    ("img", "fig2_pipeline.png", 6.6)]

C["4.2.1"] = [("p",
    "Each grayscale radiograph is converted to three identical channels and resized to "
    "224 x 224 pixels to match MobileNetV2's expected input. Pixel intensities are scaled "
    "to [0, 1] and then standardised using the ImageNet channel means "
    "(0.485, 0.456, 0.406) and standard deviations (0.229, 0.224, 0.225), so the input "
    "statistics align with those the backbone was originally trained on.")]
C["4.2.2"] = [("p",
    "No data augmentation was applied in the reported configuration. Because the "
    "convolutional backbone is frozen and its features are pre-computed and cached once, "
    "applying random augmentations per epoch would require re-running the backbone every "
    "epoch and would defeat the CPU-efficiency goal. Augmentation (e.g. small rotations "
    "and horizontal flips) is identified in Section 8.2 as a clear avenue for reducing the "
    "overfitting observed in training.")]
C["4.2.3"] = [("p",
    "No additional preprocessing (such as histogram equalisation or lung-field "
    "segmentation) was used, in order to keep the baseline simple and the comparison with "
    "the pretrained feature space clean.")]

C["4.3"] = [("p",
    "Feature extraction is performed by a deep CNN rather than hand-crafted descriptors "
    "(such as SIFT or HOG). The MobileNetV2 backbone [5], pretrained on ImageNet [7], is "
    "truncated after its final convolutional block and followed by global average pooling, "
    "yielding a 1280-dimensional embedding per image. All backbone weights are frozen. "
    "Embeddings for the train, validation and test splits are computed once and stored as "
    "NumPy arrays, so subsequent experiments reuse them instantly.")]

C["4.4"] = [("p",
    "The trainable component is a compact multi-layer perceptron head (Figure 3) that maps "
    "the 1280-dimensional feature vector to 14 logits: Linear(1280 -> 512) -> "
    "BatchNorm1d -> ReLU -> Dropout(p=0.3) -> Linear(512 -> 14). A sigmoid is applied at "
    "inference to obtain an independent probability per condition. Only this head (roughly "
    "0.66 million parameters) is optimised; the backbone is fixed. This is a deliberate "
    "transfer-learning choice that keeps training tractable on a CPU."),
    ("img", "fig3_model.png", 3.6)]

C["5.2"] = [("p",
    "All experiments were run locally on a CPU-only Windows 11 machine (no GPU / CUDA). "
    "Feature extraction over all 5,606 images completed in approximately five minutes, and "
    "training the classifier head on the cached features completed in under one minute, "
    "confirming the practicality of the frozen-feature approach on commodity hardware.")]

C["6.1"] = [("p",
    "Because the task is multi-label and highly imbalanced, plain accuracy is misleading: "
    "a trivial model that predicts 'no finding' for every image already achieves very high "
    "per-label accuracy. We therefore report the Receiver Operating Characteristic Area "
    "Under the Curve (ROC-AUC), which is threshold-independent and measures ranking quality "
    "per class, as the primary metric, alongside Precision, Recall and F1-score at a 0.5 "
    "threshold. AUC and F1 are averaged across the 14 classes (macro-averaging) so that "
    "rare findings count as much as common ones.")]

C["6.2"] = [("p",
    f"Table 3 summarises overall test-set performance against a naive all-negative "
    f"baseline. The model reaches a macro ROC-AUC of {auc:.3f} (best validation "
    f"{best_val_auc:.3f}). The accuracy column illustrates the 'accuracy paradox': the "
    f"all-negative baseline scores ~{baseline_acc*100:.0f}% per-label accuracy yet is "
    f"clinically useless (zero recall), whereas our model deliberately trades some accuracy "
    f"for substantially higher recall and a non-zero F1. Per-class results are shown in "
    f"Table 4 and Figure 7.")]

C["6.3"] = [("p",
    f"Figures 4 and 5 plot loss and validation macro-AUC against epoch. Training loss "
    f"decreases steadily, but validation loss begins rising after only a few epochs while "
    f"validation AUC peaks at epoch {best_ep} and then declines — a clear signature of "
    f"overfitting. Best-checkpoint selection on validation AUC mitigates this by saving the "
    f"model from epoch {best_ep} for final testing."),
    ("img", "fig4_loss.png", 5.2),
    ("img", "fig5_auc.png", 5.2)]

C["6.4"] = [("p",
    "Figure 6 shows row-normalised confusion matrices for four representative classes. "
    "Because the class-weighted loss emphasises the rare positive class, the model tends to "
    "achieve high recall (most true positives are caught) at the cost of many false "
    "positives, i.e. lower precision — visible as substantial mass in the predicted-positive "
    "column."),
    ("img", "fig6_confusion.png", 5.4)]

C["6.5"] = [("p",
    "Figure 8 presents example predictions: the top row shows confident, correct detections "
    "while the bottom row shows characteristic errors. Each panel lists the model's top "
    "predicted findings with confidence scores alongside the ground-truth labels."),
    ("img", "fig8_predictions.png", 6.6)]

C["6.6"] = [("p",
    f"Three patterns dominate the errors. First, performance tracks class frequency and "
    f"signal strength: findings with a distinctive global appearance such as "
    f"{top_classes[0]} (AUC {pc[top_classes[0]]['auc']:.2f}) and Effusion "
    f"(AUC {pc['Effusion']['auc']:.2f}) are detected well, whereas small, localised findings "
    f"such as Nodule (AUC {pc['Nodule']['auc']:.2f}) and the rare, heterogeneous Fibrosis "
    f"(AUC {pc['Fibrosis']['auc']:.2f}) are near or below chance. Second, precision is low "
    f"across the board because the class-weighted loss biases the model toward predicting "
    f"positives, producing many false alarms. Third, extremely rare classes such as Hernia "
    f"({counts['Hernia']} positives total) have too few examples for stable learning. A "
    f"frozen ImageNet backbone, which never saw radiographs during pretraining, also limits "
    f"the discriminability of subtle radiological texture.")]

C["7.1"] = [("p",
    f"The core objectives were met. A working multi-label classifier for all 14 findings "
    f"was built and trained entirely on a CPU using frozen-feature transfer learning, and "
    f"class imbalance was handled via class-weighted loss and AUC/F1 evaluation. The "
    f"quantitative target — a macro ROC-AUC of at least 0.65 — was achieved "
    f"({auc:.2f}). A reproducible pipeline and an interactive Streamlit demo were also "
    f"delivered. The objective of high precision was only partially met, as discussed below.")]
C["7.2"] = [("p",
    "The principal strength is efficiency: by caching features from a frozen backbone, the "
    "entire system trains in minutes on a CPU, making it highly reproducible and accessible. "
    "The design also follows sound methodology — patient-level splitting prevents leakage, "
    "class weighting counters imbalance, and threshold-independent metrics give an honest "
    "picture of performance.")]
C["7.3"] = [("p",
    f"The main weaknesses are overfitting (validation performance degrades after epoch "
    f"{best_ep}) and low precision at the default threshold, leading to a modest macro-F1 "
    f"of {macro_f1*100:.1f}%. Using only the 5,606-image sample with a frozen, "
    f"non-medical backbone caps achievable accuracy, and the noisy weak labels add further "
    f"ceiling effects.")]
C["7.4"] = [("p",
    f"On the full 112,000-image dataset, end-to-end models such as CheXNet [2] report "
    f"per-class AUCs around 0.80-0.85, and CheXpert benchmarks [3] are similar. Our macro "
    f"AUC of {auc:.2f} is lower, which is expected given that we use roughly 5% of the data, "
    f"a frozen lightweight backbone, and CPU-only training. The gap quantifies the cost of "
    f"the low-resource constraint rather than a flaw in the method, and several classes "
    f"(e.g. {top_classes[0]}, Effusion) already approach the literature range.")]

C["8.1"] = [("p",
    f"This project demonstrated that multi-label chest X-ray classification is achievable "
    f"without specialised hardware by combining a frozen ImageNet-pretrained MobileNetV2 "
    f"feature extractor with a lightweight trainable head. Trained on the 5,606-image "
    f"ChestX-ray14 sample, the system reached a macro ROC-AUC of {auc:.2f}, performing best "
    f"on findings with strong global signatures. The experiments also clearly exposed the "
    f"two central challenges of this domain — class imbalance and overfitting on limited "
    f"data — and showed why ranking metrics are essential for honest evaluation. The result "
    f"is a reproducible, accessible baseline and a working demonstration application.")]
C["8.2"] = [
    ("b", "Apply on-the-fly data augmentation (small rotations, horizontal flips, random crops) and stronger regularisation to curb the overfitting observed after the first few epochs."),
    ("b", "Fine-tune the upper backbone blocks (or adopt a chest-X-ray-pretrained backbone) instead of fully freezing them, to better capture subtle radiological texture."),
    ("b", "Tune per-class decision thresholds and train on the full 112,000-image dataset to raise precision and overall AUC toward the published CheXNet/CheXpert range."),
]

FIG_DELETE_PREFIX = "[insert"  # leftover figure placeholders to remove

# ================================================================ pass 1: map headings
heading_paras = {}      # key -> paragraph
abstract_para = None
keywords_para = None
seed_para = None
folder_para = None
objective_paras = []
future_paras = []
for p in doc.paragraphs:
    txt = p.text.strip()
    if p.style.name.startswith("Heading"):
        k = key_of(txt)
        if k:
            heading_paras[k] = p
            if k in THEME_TITLES:
                p.text = THEME_TITLES[k]
    elif txt == "Abstract":
        abstract_para = p
    elif txt.startswith("Keywords:"):
        keywords_para = p
    elif txt.startswith("Random seed:"):
        seed_para = p
    elif txt.startswith("ProjectName/"):
        folder_para = p
    elif txt.startswith("To develop") or txt.startswith("To evaluate") or txt.startswith("To achieve"):
        objective_paras.append(p)
    elif txt.startswith("Future direction"):
        future_paras.append(p)

# ================================================================ pass 2: insert section content
if abstract_para is not None:
    add_after(abstract_para._p, C["Abstract"])

for k, items in C.items():
    if k == "Abstract":
        continue
    if k in heading_paras:
        add_after(heading_paras[k]._p, items)

# ================================================================ pass 3: replace special paragraphs
if keywords_para is not None:
    keywords_para.text = ("Keywords: Chest X-ray, Multi-Label Classification, Transfer "
                          "Learning, Convolutional Neural Network, Medical Imaging, MobileNetV2")
if seed_para is not None:
    seed_para.text = (f"Random seed: {config.RANDOM_SEED}. All experiments can be reproduced "
                      f"by running download_data.py, prepare_data.py, extract_features.py and "
                      f"train.py in order (see README.md).")

# objectives (1.3) — overwrite the three template bullets, add extras after the last
obj_texts = [it[1] for it in C["1.3"]]
for i, para in enumerate(objective_paras):
    if i < len(obj_texts):
        para.text = obj_texts[i]
if len(obj_texts) > len(objective_paras) and objective_paras:
    extra = [("b", t) for t in obj_texts[len(objective_paras):]]
    add_after(objective_paras[-1]._p, extra)

# future work (8.2)
fut_texts = [it[1] for it in C["8.2"]]
for i, para in enumerate(future_paras):
    if i < len(fut_texts):
        para.text = fut_texts[i]

# folder structure (5.3)
if folder_para is not None:
    folder_para.text = (
        "chest-xray-multilabel-classifier/\n"
        "  config.py            - central settings (labels, backbone, hyper-parameters)\n"
        "  model.py             - frozen backbone + trainable classifier head\n"
        "  dataset.py           - image preprocessing and multi-label parsing\n"
        "  download_data.py     - downloads the dataset via the Kaggle API\n"
        "  prepare_data.py      - builds the patient-level train/val/test split\n"
        "  extract_features.py  - caches frozen-backbone feature vectors\n"
        "  train.py             - trains the head, evaluates, saves metrics/model\n"
        "  make_figures.py      - generates all report figures\n"
        "  app.py               - Streamlit demonstration app\n"
        "  requirements.txt     - dependencies\n"
        "  README.md            - setup and run instructions\n"
        "  artifacts/           - cached features, trained model, metrics\n"
        "  results/             - generated figures")

# ================================================================ pass 4: tables
t0, t1, t2, t3 = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]
fill_table(t0, [
    (0, f"{n_total:,}"),
    (1, "14 thoracic findings (multi-label) + 'No Finding'"),
    (2, ", ".join(labels)),
    (3, "1024 x 1024 (original); resized to 224 x 224"),
    (4, "PNG (grayscale)"),
    (5, f"{n_train} / {n_val} / {n_test}  (approx. 70% / 15% / 15%), split by patient ID to prevent leakage"),
    (6, f"Highly imbalanced (Infiltration {counts['Infiltration']} ... Hernia {counts['Hernia']} positives)"),
])
fill_table(t1, [
    (0, "PyTorch 2.7.1 (CPU build) + TorchVision 0.22.1"),
    (1, "Adam (beta1=0.9, beta2=0.999)"),
    (2, "0.001 (weight decay 1e-5)"),
    (3, "Binary Cross-Entropy with logits (per-class pos_weight for imbalance)"),
    (4, "128 (head training); 32 (feature extraction)"),
    (5, f"30 (best checkpoint at epoch {best_ep} by validation macro-AUC)"),
    (6, "Best-checkpoint selection on validation macro-AUC"),
    (7, "CPU only (no GPU)"),
    (8, "Head < 1 min on cached features; feature extraction ~5 min (one-time)"),
])
fill_table(t2, [
    (0, "Python 3.12"),
    (1, "PyTorch 2.7.1 + TorchVision 0.22.1"),
    (2, "scikit-learn, NumPy, pandas, Pillow, Matplotlib, Streamlit, kagglehub"),
    (3, "Visual Studio Code"),
])
# Table 3 (results): rows 1-4, cols My Model | Baseline | Improvement
res_rows = [
    (1, f"{per_label_acc*100:.1f} (per-label)", f"{baseline_acc*100:.1f} (all-negative)", f"{(per_label_acc-baseline_acc)*100:+.1f}"),
    (2, f"{macro_p*100:.1f}", "0.0", f"{macro_p*100:+.1f}"),
    (3, f"{macro_r*100:.1f}", "0.0", f"{macro_r*100:+.1f}"),
    (4, f"{macro_f1*100:.1f}", "0.0", f"{macro_f1*100:+.1f}"),
]
for r, a, b, c in res_rows:
    t3.rows[r].cells[1].text = a
    t3.rows[r].cells[2].text = b
    t3.rows[r].cells[3].text = c

# ---- per-class table (Table 4) + Figure 7, inserted after Table 3 ----
pcl_tbl = doc.add_table(rows=1 + config.NUM_CLASSES, cols=5)
try:
    pcl_tbl.style = "Table Grid"
except Exception:
    pass
hdr = ["Class", "ROC-AUC", "Precision", "Recall", "F1-score"]
for c, h in enumerate(hdr):
    pcl_tbl.rows[0].cells[c].text = h
for i, name in enumerate(sorted(pc, key=lambda n: -pc[n]["auc"]), start=1):
    d = pc[name]
    vals = [name, f"{d['auc']:.3f}", f"{d['precision']:.3f}", f"{d['recall']:.3f}", f"{d['f1']:.3f}"]
    for c, v in enumerate(vals):
        pcl_tbl.rows[i].cells[c].text = v

cap4 = doc.add_paragraph()
rc = cap4.add_run("Table 4: Per-class performance on the test set (sorted by ROC-AUC).")
rc.italic = True; rc.font.size = Pt(9)
figp = doc.add_paragraph(); figp.alignment = WD_ALIGN_PARAGRAPH.CENTER
figp.add_run().add_picture(str(RESULTS / "fig7_per_class_auc.png"), width=Inches(5.2))
# move all three to just after Table 3
anchor = t3._tbl
for el in (cap4._p, pcl_tbl._tbl, figp._p):
    anchor.addnext(el)
    anchor = el

# ================================================================ pass 5: cleanup placeholders/instructions
for p in list(doc.paragraphs):
    t = p.text.strip()
    low = t.lower()
    if ("INSTRUCTION" in t or low.startswith("[insert") or low.startswith("[write")
            or low.startswith("[describe") or low.startswith("[noise")):
        p._element.getparent().remove(p._element)

# ================================================================ title block at top
first = doc.paragraphs[0]
for line, size, bold in [
    (UNIV, 12, False),
    (f"Course: {COURSE}", 12, False),
    (f"Instructor: {INSTRUCTOR}", 12, False),
    (f"Author: {AUTHOR}    |    Student ID: {SID}", 12, False),
    (f"Date: {DATE}", 12, False),
    ("", 6, False),
    ("Multi-Label Chest X-ray Disease Classification using Transfer Learning", 18, True),
]:
    np_ = first.insert_paragraph_before(line)
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if np_.runs:
        np_.runs[0].bold = bold
        np_.runs[0].font.size = Pt(size)

# ================================================================ references
doc.add_paragraph()
h = doc.add_paragraph("References")
try:
    h.style = doc.styles["Heading 1"]
except KeyError:
    for kk in ("8.", "7.", "1."):          # reuse a real H1 style object from the template
        if kk in heading_paras:
            h.style = heading_paras[kk].style
            break
refs = [
    "X. Wang, Y. Peng, L. Lu, Z. Lu, M. Bagheri, and R. M. Summers, \"ChestX-ray8: Hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases,\" in Proc. IEEE CVPR, 2017, pp. 2097-2106.",
    "P. Rajpurkar et al., \"CheXNet: Radiologist-level pneumonia detection on chest X-rays with deep learning,\" arXiv:1711.05225, 2017.",
    "J. Irvin et al., \"CheXpert: A large chest radiograph dataset with uncertainty labels and expert comparison,\" in Proc. AAAI, 2019, pp. 590-597.",
    "G. Huang, Z. Liu, L. van der Maaten, and K. Q. Weinberger, \"Densely connected convolutional networks,\" in Proc. IEEE CVPR, 2017, pp. 4700-4708.",
    "M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen, \"MobileNetV2: Inverted residuals and linear bottlenecks,\" in Proc. IEEE CVPR, 2018, pp. 4510-4520.",
    "K. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning for image recognition,\" in Proc. IEEE CVPR, 2016, pp. 770-778.",
    "O. Russakovsky et al., \"ImageNet large scale visual recognition challenge,\" Int. J. Comput. Vis., vol. 115, no. 3, pp. 211-252, 2015.",
    "S. J. Pan and Q. Yang, \"A survey on transfer learning,\" IEEE Trans. Knowl. Data Eng., vol. 22, no. 10, pp. 1345-1359, 2010.",
    "C. Tan, F. Sun, T. Kong, W. Zhang, C. Yang, and C. Liu, \"A survey on deep transfer learning,\" in Proc. ICANN, 2018, pp. 270-279.",
    "E. Calli, E. Sogancioglu, B. van Ginneken, K. G. van Leeuwen, and K. Murphy, \"Deep learning for chest X-ray analysis: A survey,\" Med. Image Anal., vol. 72, 102125, 2021.",
    "A. Paszke et al., \"PyTorch: An imperative style, high-performance deep learning library,\" in Proc. NeurIPS, 2019, pp. 8024-8035.",
    "J. Deng, W. Dong, R. Socher, L.-J. Li, K. Li, and L. Fei-Fei, \"ImageNet: A large-scale hierarchical image database,\" in Proc. IEEE CVPR, 2009, pp. 248-255.",
]
for i, r in enumerate(refs, 1):
    doc.add_paragraph(f"[{i}] {r}")

doc.save(OUTPUT)
print(f"Saved report -> {OUTPUT}")
print(f"Sections filled, 4 tables + per-class table, 8 figures, 12 references.")
